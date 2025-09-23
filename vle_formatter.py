import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

st.title("📄 Excel → Word (TO first, then FROM) — Dynamic Formatting")

# ---------- Helpers ----------
def apply_case(text: str, mode: str) -> str:
    s = "" if text is None else str(text)
    if mode == "UPPERCASE":
        return s.upper()
    if mode == "lowercase":
        return s.lower()
    if mode == "Proper Case":
        return s.title()
    return s  # Original

def clean_value(val) -> str:
    if val is None:
        return ""
    try:
        if isinstance(val, float) and float(val).is_integer():
            return str(int(val))
    except Exception:
        pass
    return str(val).strip()

def add_line(container, text: str, font_size_pt: int, bold: bool = False):
    """Add a paragraph line inside a Document or a Table Cell"""
    if hasattr(container, "add_paragraph"):  # works for doc or cell
        p = container.add_paragraph()
    else:
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    return p

# ---------- App ----------
uploaded_file = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file, dtype=str)
    st.success(f"Loaded {df.shape[0]} rows × {df.shape[1]} columns")
    st.expander("Preview first rows (raw)").dataframe(df.head())

    # --- Column selection ---
    st.subheader("TO Section (appears first)")
    to_columns = st.multiselect("Choose TO fields (order matters)", df.columns.tolist())

    st.subheader("FROM Section (appears below TO)")
    from_column = st.selectbox("Choose FROM field (single column)", ["(none)"] + df.columns.tolist())
    use_from = from_column != "(none)"

    # --- Renaming TO fields ---
    st.subheader("Rename TO Fields (labels shown in output)")
    column_rename_map = {}
    for col in to_columns:
        column_rename_map[col] = st.text_input(f"Rename '{col}' →", value=col, key=f"rename_{col}")

    # --- Text casing ---
    st.subheader("Text Casing")
    case_option = st.selectbox("Apply casing to labels & values", ["Original", "UPPERCASE", "Proper Case", "lowercase"])

    # --- Font sizes ---
    st.subheader("Font Sizes")
    to_label_size = st.slider("TO label font size", 8, 36, 14)
    to_data_size  = st.slider("TO lines font size", 8, 36, 14)
    from_label_size = st.slider("FROM label font size", 8, 36, 12)
    from_data_size  = st.slider("FROM line font size", 8, 36, 12)

    # --- Bold controls ---
    st.subheader("Bold Settings")
    bold_to_label   = st.checkbox("Make 'TO:' label bold", value=True)
    bold_from_label = st.checkbox("Make 'FROM:' label bold", value=True)
    to_labels_list = [column_rename_map[c] for c in to_columns]
    bold_to_fields = st.multiselect("Make these TO lines bold", to_labels_list)
    blankline_after = st.multiselect("Add a blank line after these TO lines", to_labels_list)
    bold_from_line = st.checkbox("Make FROM address line bold", value=False)

    # --- Layout choice ---
    st.subheader("Page Layout")
    layout_choice = st.radio("Samples per page", ["1 per page", "4 per page (2×2)", "8 per page (2×4)"])

    # ---------- Preview ----------
    st.subheader("👀 Live Preview (first 2 records)")
    preview_rows = min(2, len(df))
    for i in range(preview_rows):
        row = df.iloc[i]
        st.markdown(f"<div style='font-size:{to_label_size}px; font-weight:{'700' if bold_to_label else '400'}'>{apply_case('TO:', case_option)}</div>", unsafe_allow_html=True)
        for col in to_columns:
            raw_label = column_rename_map[col]
            show_label = apply_case(raw_label, case_option)
            value = apply_case(clean_value(row.get(col, "")), case_option)
            bold_line = raw_label in bold_to_fields
            st.markdown(f"<div style='font-size:{to_data_size}px; font-weight:{'700' if bold_line else '400'}'>{show_label}: {value}</div>", unsafe_allow_html=True)
            if raw_label in blankline_after:
                st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        if use_from:
            st.markdown(f"<div style='margin-top:8px; font-size:{from_label_size}px; font-weight:{'700' if bold_from_label else '400'}'>{apply_case('FROM:', case_option)}</div>", unsafe_allow_html=True)
            from_value = apply_case(clean_value(row.get(from_column, "")), case_option)
            st.markdown(f"<div style='font-size:{from_data_size}px; font-weight:{'700' if bold_from_line else '400'}'>{from_value}</div>", unsafe_allow_html=True)
        st.markdown("<hr>", unsafe_allow_html=True)

    # ---------- DOCX Builder ----------
    def render_sample(container, row):
        add_line(container, apply_case("TO:", case_option), to_label_size, bold=bold_to_label)
        for col in to_columns:
            raw_label = column_rename_map[col]
            show_label = apply_case(raw_label, case_option)
            value = apply_case(clean_value(row.get(col, "")), case_option)
            make_bold = raw_label in bold_to_fields
            add_line(container, f"{show_label}: {value}", to_data_size, bold=make_bold)
            if raw_label in blankline_after:
                container.add_paragraph()
        if use_from:
            add_line(container, apply_case("FROM:", case_option), from_label_size, bold=bold_from_label)
            from_value = apply_case(clean_value(row.get(from_column, "")), case_option)
            add_line(container, from_value, from_data_size, bold=bold_from_line)

    def build_doc(full_df: pd.DataFrame) -> BytesIO:
        doc = Document()

        if layout_choice == "1 per page":
            for _, row in full_df.iterrows():
                render_sample(doc, row)
                doc.add_page_break()

        else:
            # decide table size
            if layout_choice == "4 per page (2×2)":
                rows_per_page, cols_per_page = 2, 2
            else:  # 8 per page
                rows_per_page, cols_per_page = 4, 2

            table = None
            count = 0
            for _, row in full_df.iterrows():
                if count % (rows_per_page * cols_per_page) == 0:
                    if count > 0:
                        doc.add_page_break()
                    table = doc.add_table(rows=rows_per_page, cols=cols_per_page)
                    for r in table.rows:
                        for c in r.cells:
                            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                            c.text = ""  # clear default

                r = (count // cols_per_page) % rows_per_page
                c = count % cols_per_page
                cell = table.cell(r, c)
                render_sample(cell, row)
                count += 1

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ---------- Download ----------
    if st.button("Generate Word (.docx)"):
        output = build_doc(df)
        st.success("✅ DOCX generated")
        st.download_button(
            "📥 Download DOCX",
            data=output,
            file_name="output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
